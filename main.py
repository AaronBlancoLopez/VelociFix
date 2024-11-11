#!/usr/bin/env python3

import argparse
import time
import re
import requests
import json
import velociraptorQueryManager as vqm

from termcolor import colored
from simple_term_menu import TerminalMenu
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

severity_ranking = {
    'HIGH': 3,
    'MEDIUM': 2,
    'LOW': 1,
}

def art():
    print("""                                                          
@@@  @@@  @@@@@@@@  @@@        @@@@@@    @@@@@@@  @@@  @@@@@@@@  @@@  @@@  @@@  
@@@  @@@  @@@@@@@@  @@@       @@@@@@@@  @@@@@@@@  @@@  @@@@@@@@  @@@  @@@  @@@  
@@!  @@@  @@!       @@!       @@!  @@@  !@@       @@!  @@!       @@!  @@!  !@@  
!@!  @!@  !@!       !@!       !@!  @!@  !@!       !@!  !@!       !@!  !@!  @!!  
@!@  !@!  @!!!:!    @!!       @!@  !@!  !@!       !!@  @!!!:!    !!@   !@@!@!   
!@!  !!!  !!!!!:    !!!       !@!  !!!  !!!       !!!  !!!!!:    !!!    @!!!    
:!:  !!:  !!:       !!:       !!:  !!!  :!!       !!:  !!:       !!:   !: :!!   
 ::!!:!   :!:        :!:      :!:  !:!  :!:       :!:  :!:       :!:  :!:  !:!  
  ::::     :: ::::   :: ::::  ::::: ::   ::: :::   ::   ::        ::   ::  :::  
   :      : :: ::   : :: : :   : :  :    :: :: :  :     :        :     :   ::
    """)
    print("\n\n\nPlease understand that this is a proof of concept intended to test the extensibility of the Velociraptor SIEM.")
    print("Code might contain bugs.")
    input("\n\n\nPress any key to continue...")


# allows to change the color a row to red if the app is vulnerable
def changeRowColor(doc, field_name, color):
    for table in doc.tables:
        for row in table.rows:
            if any(field_name in cell.text for cell in row.cells):
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = color


# writes the apps information to a .docx file
def appsInfoToDOCX(doc, data):
    doc.add_page_break()
    doc.add_heading('Software Information', level=2)
    doc.add_paragraph("\n")
    # Add a table to the document
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Publisher'
    hdr_cells[1].text = 'Display Name'
    hdr_cells[2].text = 'Display Version'
    # Formatting header
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.alignment 
            paragraph.runs[0].bold = True
            paragraph.runs[0].font.size = Pt(12)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Populate the table
    for sublist in data:
        for item in sublist:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get('Publisher') or 'N/A'
            row_cells[1].text = item.get('DisplayName') or 'N/A'
            row_cells[2].text = item.get('DisplayVersion') or 'N/A'
            # Center the text in each cell
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# writes client information to the .docx report
def clientInfoToDOCX(doc, clientInfo):
    for item in clientInfo:
        for key, value in item.items():
            doc.add_heading(key.replace('_', ' ').title(), level=2)
            if isinstance(value, dict):
                for sub_key, sub_value in value.items():
                    if isinstance(sub_value, list):
                        doc.add_paragraph(f"{sub_key.replace('_', ' ').title()}:")
                        for sub_item in sub_value:
                            doc.add_paragraph(f"  - {sub_item}", style='List Bullet')
                    else:
                        doc.add_paragraph(f"{sub_key.replace('_', ' ').title()}: {sub_value}")
            else:
                doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
            doc.add_paragraph()


def set_bold(run, bold=True):
    run.bold = bold

def set_color(run, color):
    run.font.color.rgb = color

# writes the vulnerabilities information to the .docx report
def vulnerabilitiesToDOCX(doc, vulnerabilities, app, version):
    doc.add_page_break()
    heading = doc.add_heading(f"Vulnerabilities for {app} {version}", level=1)
    for run in heading.runs:
        set_bold(run)
    for vuln in vulnerabilities:
        heading = doc.add_heading(str(vuln['CVE ID']), level=2)
        for run in heading.runs:
            set_bold(run)
        # Base severity
        severity_para = doc.add_paragraph()
        severity_run = severity_para.add_run("Base severity: ")
        set_bold(severity_run)
        value_run = severity_para.add_run(str(vuln['Base severity']))
        if str(vuln['Base severity']).upper() == "HIGH":
            set_color(value_run, RGBColor(255, 0, 0)) 
        elif str(vuln['Base severity']).upper() == "MEDIUM":
            set_color(value_run, RGBColor(255, 165, 0))
        elif str(vuln['Base severity']).upper() == "LOW":
            set_color(value_run, RGBColor(255, 255, 0))
        # Description
        description_para = doc.add_paragraph()
        description_run = description_para.add_run("Description: ")
        set_bold(description_run)
        description_para.add_run(str(vuln['Description']))
        # Exploitability score
        exploitability_para = doc.add_paragraph()
        exploitability_run = exploitability_para.add_run("Exploitability score: ")
        set_bold(exploitability_run)
        exploitability_para.add_run(str(vuln['Exploitability score']))
        # Impact score
        impact_para = doc.add_paragraph()
        impact_run = impact_para.add_run("Impact score: ")
        set_bold(impact_run)
        impact_para.add_run(str(vuln['Impact score']))
        # Weaknesses
        weaknesses_para = doc.add_paragraph()
        weaknesses_run = weaknesses_para.add_run("Weaknesses: ")
        set_bold(weaknesses_run)
        weaknesses_para.add_run(str(vuln['Weaknesses']))
        # Link
        link_para = doc.add_paragraph()
        link_run = link_para.add_run("Link: ")
        set_bold(link_run)
        link_para.add_run(str(vuln['Link']))
        doc.add_paragraph()  


# orders vulnerabilities by severity
def order_vulnerabilities(vulnerabilities):
    # Convert severity to uppercase to handle any case differences
    return sorted(vulnerabilities, key=lambda vuln: severity_ranking.get(vuln['Base severity'].upper(), 0), reverse=True)


# checks if the first two words of a string are alphabetic
def first_two_words(string):
    elements = string.split()
    if len(elements) >= 2 and elements[0].isalpha() and elements[1].isalpha():
        return 1
    return 0

# retrieves the possible CPEs for a given app name and version
def find_cpes(name, version):
    base_url = "https://nvd.nist.gov/products/cpe/search/results"
    params = {
        "namingFormat": "2.3",
        "keyword": f"{name} {version}"
    }
    response = requests.get(base_url, params=params)
    content = response.text
    cpe_matches = re.findall(r'cpe:(.*?)<', content)
    return cpe_matches[0] if cpe_matches else None


# retrieves the CVE details for a given CPE
def fetch_cve_details(cpe_string):
    base_url = "https://services.nvd.nist.gov/rest/json/cves/2.0"
    url = f"{base_url}?cpeName=cpe:{cpe_string}" 
    # NVD API rate limit is 5 requests every 30 seconds (50 if api key is given)
    time.sleep(6)
    response = requests.get(url)
    if response.status_code != 200:
        print(colored(f"Error: Unable to retrieve CVE data for CPE: {cpe_string}. Status code: {response.status_code}", "red"))
        return []
    try:
        data = response.json()
    except json.JSONDecodeError:
        print(colored(f"Error decoding JSON for CPE: {cpe_string}. Skipping.", "red"))
        return []
    if data["vulnerabilities"]:
        all_cve_details = []
        for cve_item in data["vulnerabilities"]:
            cve_id = cve_item["cve"]["id"]
            description_text = cve_item["cve"]["descriptions"][0]["value"]
            link = f"https://nvd.nist.gov/vuln/detail/{cve_id}"
            weaknesses = []
            for problem_type in cve_item["cve"]["weaknesses"]:
                for description in problem_type["description"]:
                    weaknesses.append(description["value"])
            if cve_item.get("cve", {}).get("metrics", {}).get("cvssMetricV31") is not None:
                baseSeverity = cve_item["cve"]["metrics"]["cvssMetricV31"][0]["cvssData"]["baseSeverity"]    
                exploitabilityScore = cve_item["cve"]["metrics"]["cvssMetricV31"][0]["exploitabilityScore"]
                impactScore = cve_item["cve"]["metrics"]["cvssMetricV31"][0]["impactScore"]
            elif cve_item.get("cve", {}).get("metrics", {}).get("cvssMetricV2") is not None:
                baseSeverity = cve_item["cve"]["metrics"]["cvssMetricV2"][0]["baseSeverity"]
                exploitabilityScore = cve_item["cve"]["metrics"]["cvssMetricV2"][0]["exploitabilityScore"]
                impactScore = cve_item["cve"]["metrics"]["cvssMetricV2"][0]["impactScore"]
            all_cve_details.append({
                "CVE ID": cve_id,
                "Description": description_text,
                "Weaknesses": ", ".join(weaknesses),
                "Base severity": baseSeverity,
                "Exploitability score": exploitabilityScore,
                "Impact score": impactScore,
                "Link": link,
            })
        return all_cve_details
    return []


def main():
    
    parser = argparse.ArgumentParser()
    parser.add_argument('--config', type=str,
                        help='Path to the api_client configuration file. You can generate such file with "velociraptor config api_client"')
    parser.add_argument('--repository', type=str,
                        help='URL of the repository where the apps are stored. Do not include the last backslash ("/"). This is an optional parameter, but you will only get the vulnerability detection report, as patching can not be done without the repository.')
    args = parser.parse_args()

    art()

    # main menu configuration
    main_menu_title = "  Main Menu.\n  Press Q or Esc to quit. \n"
    main_menu_items = ["See connected clients", "Scan client", "Quit"]
    main_menu_cursor = "> "
    main_menu_cursor_style = ("fg_red", "bold")
    main_menu_style = ("bg_red", "fg_yellow")
    main_menu_exit = False
    # main menu
    main_menu = TerminalMenu(
        menu_entries=main_menu_items,
        title=main_menu_title,
        menu_cursor=main_menu_cursor,
        menu_cursor_style=main_menu_cursor_style,
        menu_highlight_style=main_menu_style,
        cycle_cursor=True,
        clear_screen=True,
    )
    # main menu loop
    while not main_menu_exit:
        main_sel = main_menu.show()


        # FIRST OPTION: lists all currently connected clients
        if main_sel == 0:
            print(" List of currently connected clients:\n")
            for client in vqm.getClients(args.config):
                print("- " + client)
            input("\n\nPress Enter to go back to the main menu...")



        # SECOND OPTION: scan the apps installed in a client
        elif main_sel == 1:
            clients = vqm.getClients(args.config)
            # client menu configuration
            client_menu = TerminalMenu(
                menu_entries=clients,
                title="Select a client to scan",
                menu_cursor=main_menu_cursor,
                menu_cursor_style=main_menu_cursor_style,
                menu_highlight_style=main_menu_style,
                cycle_cursor=True,
                clear_screen=True,
            )
            print("Select a client to scan:\n")
            client_sel = client_menu.show()
            # .docx creation from JSON data
            doc = Document()
            doc.add_heading('Client Information Report', 0)
            # get and write client and apps information to the .docx file
            client_info = vqm.getClientInfo(args.config, clients[client_sel])
            apps = vqm.getApps(args.config, clients[client_sel])
            clientInfoToDOCX(doc, client_info)
            appsInfoToDOCX(doc, apps)
            # remember 32 and 64 bit apps are separated
            apps_32 = apps[0]
            apps_64 = apps[1]
            print(" List of installed applications:\n")
            # print the apps and search for possible CVEs
            max_name_length = max(len(app['DisplayName']) for app in apps_32 + apps_64)
            vulnerable = []
            # apps review loop
            for app in apps_32:
            
                if first_two_words(app['DisplayName']):
                    name = ' '.join(app['DisplayName'].split()[:2])
                else:
                    name = app['DisplayName'].split()[0]
                version = app['DisplayVersion']
                # if cpe is found for given keywords, CVEs are fetched 
                # and written to the .docx file
                if cpe := find_cpes(name, version):
                    vulnerable.append(name)
                    cves = fetch_cve_details(cpe)
                    cves_Sorted = order_vulnerabilities(cves)
                    changeRowColor(doc, app['DisplayName'], RGBColor(255, 0, 0))
                    vulnerabilitiesToDOCX(doc, cves_Sorted, app['DisplayName'], version)
                    print(colored(" - {:{}} {}".format(app['DisplayName'], max_name_length, version), "red"))
                    continue
                print(" - {:{}} {}".format(app['DisplayName'], max_name_length, version))
            

            for app in apps_64:
                name = ' '.join(app['DisplayName'].split()[:2])
                version = app['DisplayVersion']
                if cpe := find_cpes(name, version):
                    vulnerable.append(name)
                    cves = fetch_cve_details(cpe)
                    cves_Sorted = order_vulnerabilities(cves)
                    changeRowColor(doc, app['DisplayName'], RGBColor(255, 0, 0))
                    vulnerabilitiesToDOCX(doc, cves_Sorted, app['DisplayName'], version)
                    print(colored(" - {:{}} {}".format(app['DisplayName'], max_name_length, version), "red"))
                    continue
                print("- {:{}} {}".format(app['DisplayName'], max_name_length, version))

            # save the .docx file with name clientID_Information_Report.docx
            client_abreviation = clients[client_sel].split(".")[1]
            doc.save(f"{client_abreviation}_Information_Report.docx")

            # check if there are vulnerable apps and ask if the user wants to update them
            if vulnerable:
                print(colored(f"\n\n [!] {len(vulnerable)} vulnerable applications found", "red"))
                for app in vulnerable:
                    print(colored(f"    - {app}", "red"))
                option = input("\n Would you like to try to update the vulnerable applications? (y/n): ")
                if option.lower() == "y":
                    if args.repository:               
                        for app in vulnerable:
                            print(f"\n Trying to update {app}...")
                            # the standard in the repository is the app name with spaces replaced by underscores
                            # the same way the app is searched in the NVD database, but with the replacement
                            app = app.replace(" ", "_")
                            response = vqm.download(args.config, clients[client_sel], app, args.repository)
                            if(response == 1):
                                print(colored(f" [!] The file {app}.msi was not found in the given repository.\nPlease try it again later.", "red"))
                                continue
                            elif(response == -1):
                                print(colored(""))
                            print(colored(f" [+] {app} downloaded successfully!\n", "green"))
                            vqm.installation(args.config, clients[client_sel], app)
                            print(colored(f"\n [+] {app} updated successfully!", "green"))
                        print("\n You can find a report of the scan on this folder called by the name of the client's ID. ")
                        input(" Press any key to exit...")
                    else:
                        print('\033[38;2;255;165;0m' + "\n\n [!] No repository URL was provided. Vulnerable applications can not be updated." + '\033[0m')
                else:
                    print("\n You can find a report of the scan on this folder called by the name of the client's ID. ")
                    input(" Press any key to exit...")
            else:
                print(colored("\n\n [+] No vulnerable applications found", "green"))


        # exit the program
        elif main_sel == 2 or main_sel == None:
            main_menu_exit = True
            print(" Quit Selected")



if __name__ == '__main__':
    main()